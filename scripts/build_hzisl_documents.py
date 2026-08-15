from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "public" / "documents"
BRAND_DIR = ROOT / "public" / "brand"

SEASON = "2026–27"
VERSION = "1.0"
SPORT_RULES_VERSION = "1.1"
ISSUE_DATE = "15 August 2026"
REVIEW_DATE = "30 June 2027"

# compact_reference_guide preset, with a named HZISL brand-color override.
FONT = "Calibri"
NAVY = "003478"
BLUE = "2E74B5"
RED = "D71920"
INK = "0B2545"
MUTED = "526173"
PALE_BLUE = "E8EEF5"
PALE_RED = "FCEBEC"
PALE_GOLD = "FFF8E8"
PALE_GRAY = "F4F6F9"
WHITE = "FFFFFF"
LINE = "C9D2DE"
SUCCESS = "E7F4EC"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120

IFAB_URL = "https://www.theifab.com/laws-of-the-game-documents/"
FIBA_2024_URL = (
    "https://assets.fiba.basketball/image/upload/"
    "documents-corporate-fiba-official-rules-2024-v10a.pdf"
)
FIBA_RULES_HUB_URL = "https://refereeing.fiba.basketball/en/rules"
FIVB_RULES_URL = "https://www.fivb.com/wp-content/uploads/2025/01/FIVB-Volleyball_Rules2025_2028-EN.pdf"
FIVB_BASIC_RULES_URL = "https://www.fivb.com/volleyball/the-game/basic-rules/"


def _set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, width in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def _set_cell_fill(cell, fill: str | None) -> None:
    if not fill:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_table_geometry(table, widths: Sequence[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _style_table_borders(table, color: str = LINE, size: int = 6) -> None:
    edges = {
        edge: {"val": "single", "sz": size, "space": 0, "color": color}
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV")
    }
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge_name, settings in edges.items():
        edge = OxmlElement(f"w:{edge_name}")
        for key, value in settings.items():
            edge.set(qn(f"w:{key}"), str(value))
        tbl_borders.append(edge)


def _format_cell_paragraph(paragraph, *, size: float = 9.5, bold: bool = False, color: str = INK,
                           align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold, color=color)


def _set_cell_text(cell, text: str, *, size: float = 9.5, bold: bool = False,
                   color: str = INK, fill: str | None = None,
                   align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = text
    _set_cell_fill(cell, fill)
    _format_cell_paragraph(cell.paragraphs[0], size=size, bold=bold, color=color, align=align)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    _set_run_font(run, size=8.5, color=MUTED)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.extend([r_fonts, color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13.5, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        r_pr = styles[style_name]._element.get_or_add_rPr()
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:val"), "en-US")


def _add_numbering(doc: Document) -> tuple[int, int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]

    def make_numbering(abstract_id: int, num_id: int, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), NAVY)
        r_pr.append(color)
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    bullet_abs = max(abstract_ids, default=0) + 1
    decimal_abs = bullet_abs + 1
    bullet_num = max(num_ids, default=0) + 1
    decimal_num = bullet_num + 1
    make_numbering(bullet_abs, bullet_num, "bullet", "•")
    make_numbering(decimal_abs, decimal_num, "decimal", "%1.")
    return bullet_num, decimal_num, decimal_abs


def _new_num_id(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


class HZDocument:
    def __init__(
        self,
        title: str,
        short_title: str,
        logo: str = "hzisl-main.png",
        *,
        version: str = VERSION,
    ) -> None:
        self.doc = Document()
        self.title = title
        self.short_title = short_title
        self.logo_path = BRAND_DIR / logo
        self.version = version
        _configure_styles(self.doc)
        self.bullet_num, self.decimal_num, self.decimal_abstract = _add_numbering(self.doc)
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        core = self.doc.core_properties
        core.title = self.title
        core.subject = f"HZISL {SEASON} competition document"
        core.author = "Hsinchu–Zhubei International Schools League"
        core.keywords = "HZISL, interschool sports, football, basketball, volleyball"
        core.comments = f"Version {self.version}; issued {ISSUE_DATE}; review {REVIEW_DATE}."

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        if self.logo_path.exists():
            logo_run = hp.add_run()
            picture = logo_run.add_picture(str(self.logo_path), width=Inches(0.36))
            doc_pr = picture._inline.docPr
            doc_pr.set("descr", "HZISL league logo")
        text_run = hp.add_run(f"  HZISL  |  {self.short_title.upper()}  |  {SEASON}")
        _set_run_font(text_run, size=8.5, bold=True, color=MUTED)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        left = fp.add_run(f"Official competition document  •  v{self.version}  •  {ISSUE_DATE}")
        _set_run_font(left, size=8.5, color=MUTED)
        page_label = fp.add_run("\tPage ")
        _set_run_font(page_label, size=8.5, color=MUTED)
        _add_page_field(fp)

    def add_cover(self, kicker: str, subtitle: str, *, compact: bool = False) -> None:
        # Header pattern: editorial_cover; compact is a named operational override.
        if not compact:
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(14)
            if self.logo_path.exists():
                spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture = spacer.add_run().add_picture(str(self.logo_path), width=Inches(1.18))
                picture._inline.docPr.set("descr", "HZISL league logo")
        kicker_p = self.doc.add_paragraph()
        kicker_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not compact else WD_ALIGN_PARAGRAPH.LEFT
        kicker_p.paragraph_format.space_after = Pt(5)
        kr = kicker_p.add_run(kicker.upper())
        _set_run_font(kr, size=9.5, bold=True, color=RED)
        title_p = self.doc.add_paragraph(style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not compact else WD_ALIGN_PARAGRAPH.LEFT
        title_p.add_run(self.title)
        for run in title_p.runs:
            _set_run_font(run, size=28 if not compact else 25, bold=True, color=NAVY)
        subtitle_p = self.doc.add_paragraph(style="Subtitle")
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not compact else WD_ALIGN_PARAGRAPH.LEFT
        subtitle_p.add_run(subtitle)
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER if not compact else WD_ALIGN_PARAGRAPH.LEFT
        meta.paragraph_format.space_after = Pt(14)
        mr = meta.add_run(f"Season {SEASON}  •  Version {self.version}  •  Effective 5 September 2026")
        _set_run_font(mr, size=9.5, bold=True, color=MUTED)

    def add_heading(self, text: str, level: int = 1, *, page_break_before: bool = False) -> None:
        if page_break_before and self.doc.paragraphs and not self.doc.paragraphs[-1].text.strip():
            empty = self.doc.paragraphs[-1]._element
            empty.getparent().remove(empty)
        heading = self.doc.add_heading(text, level=level)
        heading.paragraph_format.page_break_before = page_break_before

    def add_para(self, text: str = "", *, bold_lead: str | None = None, italic: bool = False,
                 size: float | None = None, color: str | None = None,
                 keep_with_next: bool = False, after: float | None = None) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.keep_with_next = keep_with_next
        if after is not None:
            p.paragraph_format.space_after = Pt(after)
        if bold_lead and text.startswith(bold_lead):
            lead = p.add_run(bold_lead)
            _set_run_font(lead, size=size, bold=True, color=color or INK)
            rest = p.add_run(text[len(bold_lead):])
            _set_run_font(rest, size=size, italic=italic, color=color or INK)
        else:
            run = p.add_run(text)
            _set_run_font(run, size=size, italic=italic, color=color or INK)

    def add_bullets(self, items: Iterable[str], *, size: float = 11, after: float = 4) -> None:
        for item in items:
            p = self.doc.add_paragraph()
            _apply_numbering(p, self.bullet_num)
            p.paragraph_format.space_after = Pt(after)
            run = p.add_run(item)
            _set_run_font(run, size=size, color=INK)

    def add_numbered(self, items: Iterable[str]) -> None:
        num_id = _new_num_id(self.doc, self.decimal_abstract)
        for item in items:
            p = self.doc.add_paragraph()
            _apply_numbering(p, num_id)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(item)
            _set_run_font(run, size=11, color=INK)

    def add_callout(self, label: str, text: str, *, fill: str = PALE_BLUE,
                    accent: str = NAVY) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        _set_repeat_header(table.rows[0])
        _set_table_geometry(table, [CONTENT_WIDTH_DXA])
        _style_table_borders(table, color=accent, size=8)
        cell = table.cell(0, 0)
        _set_cell_fill(cell, fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        lead = p.add_run(f"{label.upper()}  ")
        _set_run_font(lead, size=9.5, bold=True, color=accent)
        body = p.add_run(text)
        _set_run_font(body, size=10.2, color=INK)
        self.add_para("", after=2)

    def add_table(self, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int],
                  *, body_size: float = 9.4, header_fill: str = PALE_BLUE,
                  alignments: Sequence | None = None) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        _style_table_borders(table)
        _set_repeat_header(table.rows[0])
        for idx, header in enumerate(headers):
            _set_cell_text(
                table.rows[0].cells[idx], header, size=9.2, bold=True,
                color=NAVY, fill=header_fill,
                align=(alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT),
            )
        for row_data in rows:
            row = table.add_row()
            for idx, value in enumerate(row_data):
                _set_cell_text(
                    row.cells[idx], str(value), size=body_size,
                    align=(alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT),
                )
        _set_table_geometry(table, widths)
        self.add_para("", after=2)

    def add_source(self, label: str, url: str, note: str = "") -> None:
        p = self.doc.add_paragraph()
        _apply_numbering(p, self.bullet_num)
        p.paragraph_format.space_after = Pt(4)
        lead = p.add_run(f"{label}: ")
        _set_run_font(lead, size=9.5, bold=True, color=INK)
        _add_hyperlink(p, "Official source", url)
        if note:
            tail = p.add_run(f" — {note}")
            _set_run_font(tail, size=9.5, color=MUTED)

    def save(self, filename: str) -> Path:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        target = OUTPUT_DIR / filename
        self.doc.save(target)
        return target


def add_document_control(d: HZDocument, purpose: str, audience: str) -> None:
    d.add_table(
        ["Document", "Value"],
        [
            ["Purpose", purpose],
            ["Audience", audience],
            ["Owner", "HZISL Council and League Coordinator"],
            ["Version cycle", f"Version {VERSION}; issued {ISSUE_DATE}; review by {REVIEW_DATE}"],
        ],
        [2350, 7010],
        body_size=9.6,
    )


def build_handbook() -> Path:
    d = HZDocument("Competition Handbook & Operating Statutes", "Handbook")
    d.add_cover(
        "League governance and operations",
        "The common framework for every HZISL football, basketball and volleyball division.",
    )
    d.add_callout(
        "Core promise",
        "Safe, fair and dependable interschool competition for Middle School and High School students in Hsinchu and Zhubei.",
    )
    add_document_control(
        d,
        "Set the league’s governance, eligibility, scheduling, safety, conduct and result-management rules.",
        "School leaders, athletic directors, coordinators, coaches, officials, students and families.",
    )
    d.add_heading("1. Identity, purpose and scope", page_break_before=True)
    d.add_para(
        "The Hsinchu–Zhubei International Schools League (HZISL) is a cooperative interschool sports league. "
        "It exists to provide regular, developmentally appropriate competition while protecting student welfare, school relationships and the integrity of results."
    )
    d.add_bullets([
        "Member schools for the 2026–27 season: HCAS, HIA, PAS, HIS, HAS and Korrnell Academy (KA).",
        "Sports: football, basketball and volleyball.",
        "Divisions: MS Boys, MS Girls, HS Boys and HS Girls in each sport.",
        "Season structure: a home-and-away round robin over 10 matchdays, followed by a separately published Finals Day format.",
        "Normal matchdays are Saturday mornings; MS fixtures start at 09:00 and HS fixtures at 10:30 unless the published schedule states otherwise.",
    ])
    d.add_heading("2. Values and standards")
    d.add_bullets([
        "Student safety comes before schedule, score or competitive advantage.",
        "Officials, opponents, hosts and facilities are treated with respect.",
        "Schools communicate early, accurately and through their designated representative.",
        "The signed match report is the official record; the app is the public display of verified information.",
        "Participation is inclusive and decisions are made without discrimination, consistent with applicable law and each school’s safeguarding duties.",
    ])

    d.add_heading("3. Governance and decision-making", page_break_before=True)
    d.add_heading("3.1 HZISL Council", level=2)
    d.add_para(
        "Each member school appoints one voting representative. The Council approves the season calendar, competition documents, membership and material rule changes. "
        "The League Coordinator chairs meetings, maintains records, publishes decisions and handles routine administration."
    )
    d.add_table(
        ["Decision", "Standard"],
        [
            ["Quorum", "Four of the six member schools represented."],
            ["Routine decision", "Simple majority of representatives present and voting."],
            ["Statute, membership or in-season rule change", "Two-thirds of all member schools, except an immediate safety direction."],
            ["Conflict of interest", "The affected school discloses the conflict and does not decide its own case."],
            ["Urgent interpretation", "Coordinator issues a written interim decision; Council reviews it at the next reasonable opportunity."],
        ],
        [2920, 6440],
    )
    d.add_heading("3.2 Member-school responsibilities", level=2)
    d.add_bullets([
        "Name a school representative and a matchday contact whose details are kept current.",
        "Certify eligibility, consent and medical clearance before a student participates.",
        "Provide qualified adult supervision from arrival until departure.",
        "When hosting, provide a safe venue, first-aid capability, officials and result administration.",
        "Respond promptly to league communications, discipline reviews and schedule decisions.",
    ])
    d.add_heading("4. Eligibility and rosters")
    d.add_bullets([
        "A player must be currently enrolled at the school represented and certified by that school for the declared MS or HS division.",
        "A younger student may play in an older division only when the school authorizes it and student welfare is protected; an older student may never play down.",
        "A player may not represent two schools, two teams in the same division, or both MS and HS on the same matchday without prior written HZISL approval.",
        "The school determines participation in Boys or Girls competition under its inclusion policy, applicable law and confidential student-support process; private information is not placed on public league records.",
        "The season roster and eligibility declaration must be submitted before the team’s first fixture. Matchday participants must appear on the certified roster.",
        "The school retains consent, emergency-contact and medical details. Only information necessary to run the competition is shared with HZISL."
    ])
    d.add_callout(
        "No medical data in the app",
        "The public HZISL app may show team and competition information. Completed rosters, incident reports and medical information must use the approved secure school-to-league channel.",
        fill=PALE_RED,
        accent=RED,
    )

    d.add_heading("5. Scheduling, changes and cancellations")
    d.add_numbered([
        "The published HZISL schedule is binding. A school that identifies a conflict contacts the Coordinator and opponent as early as possible.",
        "A non-emergency change should be requested at least 72 hours before kick-off. A change is not effective until the Coordinator confirms it in writing.",
        "The host monitors weather, air quality, access and facility safety. For a morning cancellation, the host should issue its recommendation by 06:30 when practicable.",
        "The referee has final authority to suspend or abandon play at the venue for unsafe conditions. The host retains its non-delegable authority to close a facility.",
        "The Coordinator decides whether an unplayed or abandoned fixture is rescheduled, replayed, recorded as a forfeit or allowed to stand, after considering safety, responsibility and the time completed.",
    ])
    d.add_heading("6. Matchday operations")
    d.add_bullets([
        "The home school is the host and follows the HZISL Host a Game Guide.",
        "The host confirms facilities, changing arrangements, water access, first aid, emergency procedures and officials before teams arrive.",
        "Visiting teams arrive at least 30 minutes before their fixture unless otherwise agreed, bring their certified roster and wear a distinguishable uniform.",
        "Each team identifies one head coach and one captain. Only authorized team personnel use the technical or bench area.",
        "The host protects a spectator area separate from teams and officials and applies its campus-access rules to all visitors.",
    ])

    d.add_heading("7. Officials and authority")
    d.add_para(
        "Match officials apply the HZISL sport rules and the incorporated international code. Their decisions on facts connected with play are final. "
        "A school may report a rules-application or administrative concern, but it may not use a protest to relitigate ordinary judgment calls."
    )
    d.add_bullets([
        "Football: the referee controls the match from pre-match inspection through submission of the report.",
        "Basketball: the referee crew controls play; the scorer and timer maintain the table record under the crew’s authority.",
        "Volleyball: the referee crew controls play; the scorer records rotations, set scores, substitutions and sanctions under the crew’s authority.",
        "Coaches communicate respectfully through the permitted channel and keep substitutes and spectators under control.",
        "Abuse, threats, discriminatory conduct or unauthorized entry onto the playing area is reported immediately and may lead to removal and further sanction.",
    ])
    d.add_heading("8. Health, safety and safeguarding")
    d.add_bullets([
        "A responsible adult from each team remains present and holds access to emergency contacts and relevant medical information.",
        "The host provides a stocked first-aid kit, a trained first-aid responder, an emergency action plan and clear access for emergency services.",
        "A player with a suspected concussion or serious injury takes no further part that day unless cleared under the school’s medical protocol by an appropriately qualified professional.",
        "No student is left alone with an unrelated adult in a closed or isolated area. Changing areas follow host safeguarding and privacy rules.",
        "Photography and publication follow school consent rules. Completed forms containing personal data are never posted publicly.",
    ])
    d.add_heading("9. Conduct and discipline", page_break_before=True)
    d.add_table(
        ["Participant", "Expected conduct"],
        [
            ["Players", "Compete honestly; respect decisions; avoid dangerous, abusive or discriminatory conduct; shake hands or acknowledge opponents after play."],
            ["Coaches", "Model calm behavior; protect students; control the bench; communicate through the proper channel; do not confront officials after the match."],
            ["Spectators", "Encourage without abuse, interference or entry into team/official areas; follow host instructions."],
            ["Officials and hosts", "Act impartially; explain administrative steps clearly; document serious incidents promptly."],
        ],
        [1800, 7560],
        body_size=9.2,
    )
    d.add_para(
        "An ejection or sending-off carries an automatic suspension from the team’s next HZISL match, pending written review. "
        "The Coordinator may impose additional measures after giving the school a fair opportunity to provide relevant information."
    )
    d.add_heading("10. Results, standings and records")
    d.add_numbered([
        "The host scorer and match official complete the Official Match Report immediately after the fixture.",
        "Both head coaches review the score and recorded sanctions and sign the form. A signature confirms the record was received; it does not waive a timely protest.",
        "The host sends the completed report to the Coordinator within 30 minutes of the final whistle/buzzer, or as soon as connectivity allows.",
        "The Coordinator verifies and publishes the result. If the app differs from the signed report, the signed report controls until formally amended.",
        "League-table methods and sport-specific forfeits are stated in the Football, Basketball and Volleyball Rules. The published app table is provisional until the season is certified."
    ])
    d.add_heading("11. Protests and review")
    d.add_bullets([
        "A coach records an intention to protest on the match report before signing when practicable.",
        "The school representative submits the written protest within 24 hours, identifying the rule or administrative error and the requested remedy.",
        "Video may be considered as supporting evidence but does not create video review during a match.",
        "A three-person neutral panel appointed from non-involved schools decides material protests. The panel may confirm, correct, replay or forfeit a result when permitted by these rules.",
        "Safeguarding matters follow the relevant school and legal reporting route immediately and are not delayed by the competition protest process."
    ])
    d.add_heading("12. Rule hierarchy and amendments")
    d.add_para("When documents appear to conflict, apply the following order:")
    d.add_numbered([
        "Applicable law and the host school’s immediate safeguarding, medical and facility-safety duties.",
        "A written HZISL season bulletin approved under these statutes.",
        "The HZISL sport-specific rules for the match concerned.",
        "This Handbook for general league operations.",
        "The incorporated IFAB, FIBA or FIVB rules, only where HZISL has not made a local modification.",
    ])
    d.add_para(
        "Material in-season changes should be avoided. A safety correction may take immediate effect; other changes apply from the date stated in the written notice."
    )
    d.add_heading("Official references")
    d.add_source("IFAB Laws of the Game 2026/27", IFAB_URL, "football baseline, accessed 15 August 2026")
    d.add_source("FIBA Official Basketball Rules 2024", FIBA_2024_URL, "basketball baseline retained for the full HZISL season unless a written bulletin says otherwise")
    d.add_source("FIVB Official Volleyball Rules 2025–2028", FIVB_RULES_URL, "volleyball baseline, subject to the published HZISL local modifications")
    return d.save("HZISL_Competition_Handbook_2026-27.docx")


def build_football_rules() -> Path:
    d = HZDocument(
        "Football Competition Rules",
        "Football Rules",
        "hzisl-football.png",
        version=SPORT_RULES_VERSION,
    )
    d.add_cover(
        "Sport regulations",
        "Applies to MS Boys, MS Girls, HS Boys and HS Girls.",
    )
    d.add_callout(
        "Match format",
        "6-a-side • unlimited return substitutions • 2 × 40-minute halves • 10-minute half-time • league matches may finish as a draw.",
    )
    d.add_heading("1. Governing code")
    d.add_para(
        "Matches are played under the IFAB Laws of the Game 2026/27, as modified below for HZISL youth competition. "
        "These local rules control where they differ from IFAB. The HZISL Competition Handbook controls league administration, eligibility, safety and protests."
    )
    d.add_heading("2. Match and team format")
    d.add_table(
        ["Item", "HZISL rule"],
        [
            ["Players", "6 on the field per team, including one goalkeeper. A team must have 6 ready to start; a match may not continue with fewer than 4."],
            ["Matchday roster", "Maximum 18 eligible players, plus authorized team staff."],
            ["Duration", "Two equal halves of 40 minutes."],
            ["Half-time", "10 minutes; changed only with the referee’s permission for safety or facility necessity."],
            ["Added time", "At the referee’s discretion for injuries, substitutions, discipline, time-wasting and other significant delay."],
            ["League tie", "No extra time; an equal score is recorded as a draw."],
            ["Knockout tie", "The separately published Finals Day regulations state the tie-breaking procedure."],
        ],
        [2050, 7310],
        body_size=9.5,
    )
    d.add_heading("3. Field, goals and ball")
    d.add_bullets([
        "The host provides a safely marked natural or artificial surface appropriate for 6-a-side play. The referee may require hazards to be removed or may refuse an unsafe field.",
        "Goals must be securely anchored. Nets, corner markings and technical areas should be provided where practicable.",
        "A size 5 match ball in good condition is used in every division. The host supplies at least two match-quality spare balls.",
        "Where local field dimensions differ, both teams are informed before matchday. The same safe field is used for comparable divisions whenever practicable.",
    ])
    d.add_heading("4. Player equipment")
    d.add_bullets([
        "Matching numbered shirts, shorts and socks are required. Goalkeepers wear colors distinguishable from both teams and the officials.",
        "Shinguards covered by socks and appropriate footwear are compulsory.",
        "Jewelry and unsafe equipment are not permitted. Medical or religious coverings must be safe and approved at the referee’s inspection.",
        "If colors conflict, the home team changes or wears bibs unless the teams agreed another solution in advance.",
    ])

    d.add_heading("5. Substitutions")
    d.add_bullets([
        "There is no limit on the number of substitutions. HZISL uses unlimited return substitutions from the named matchday roster to support safe youth participation.",
        "A substitution is made during a stoppage, at the halfway line, after the referee is informed and signals entry.",
        "A substituted player may return later. All players and substitutes remain under the referee’s authority.",
        "A goalkeeper change requires the referee’s permission and occurs during a stoppage.",
        "A suspected concussion or serious injury is managed under the school medical protocol; player welfare overrides substitution limits or competitive considerations.",
    ])
    d.add_heading("6. Start, restarts and offside")
    d.add_para(
        "Kick-offs, throw-ins, goal kicks, corner kicks, dropped balls, free kicks, penalty kicks and offside are administered under the current IFAB Laws. "
        "There is no video assistant referee or on-field video review."
    )
    d.add_heading("7. Fouls, misconduct and technical areas")
    d.add_bullets([
        "Careless, reckless and excessive-force challenges are sanctioned under IFAB Law 12. Student age does not reduce protection from dangerous play.",
        "Two cautions in the same match result in a sending-off. A sent-off player or team official leaves the playing area and may not be replaced.",
        "A sending-off triggers an automatic suspension from the next HZISL match pending league review. Additional sanction may follow for serious foul play, violent conduct, abuse or discrimination.",
        "Only authorized players and staff occupy the technical area. Coaches are responsible for substitutes and team behavior.",
        "The referee reports cautions, sendings-off, injuries and exceptional incidents on the Official Match Report or Incident Report.",
    ])
    d.add_heading("8. Time management and safety stoppages")
    d.add_bullets([
        "The referee keeps official time. A visible venue clock is informative unless the referee designates it as official.",
        "Drinks or cooling breaks may be authorized for heat and safety. They are not coaching time-outs and are added to playing time as appropriate.",
        "Lightning, severe weather, unsafe air quality, facility intrusion or medical emergency permits immediate suspension. The host emergency plan applies.",
        "The referee and host record the score, time played and reason if a match is suspended or abandoned. The Coordinator determines the competition outcome.",
    ])
    d.add_heading("9. Results and standings", page_break_before=True)
    d.add_table(
        ["Result", "League points"],
        [["Win", "3"], ["Draw", "1"], ["Loss", "0"], ["Forfeit", "0; opponent normally records a 3–0 win"]],
        [6200, 3160],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    d.add_para(
        "Teams are ranked by points, then head-to-head points among tied teams, overall goal difference, goals scored, disciplinary record and, if still required for a final placing, a neutral playoff or draw determined by the Council. "
        "The Coordinator may adjust a forfeit score when the abandoned match score was already more favorable to the non-offending team."
    )
    d.add_heading("10. Match report and protest")
    d.add_numbered([
        "The referee and host scorer confirm the final score and disciplinary record.",
        "Both coaches review and sign the Official Match Report immediately after the match.",
        "A coach notes an intention to protest before signing when practicable; the school representative submits the written protest within 24 hours.",
        "A protest may concern rule application, eligibility or administration, but not an ordinary referee judgment on facts connected with play.",
    ])
    d.add_heading("Official reference")
    d.add_source("IFAB Laws of the Game 2026/27", IFAB_URL, "including permitted youth modifications to duration, player numbers and return substitutions")
    return d.save("HZISL_Football_Rules_2026-27.docx")


def build_basketball_rules() -> Path:
    d = HZDocument(
        "Basketball Competition Rules",
        "Basketball Rules",
        "hzisl-basketball.png",
        version=SPORT_RULES_VERSION,
    )
    d.add_cover(
        "Sport regulations",
        "Applies to MS Boys, MS Girls, HS Boys and HS Girls.",
    )
    d.add_callout(
        "Game format",
        "5-a-side • unlimited substitutions • 4 × 8-minute quarters • 4-minute quarter intervals • 10-minute half-time • 3-minute overtime periods.",
    )
    d.add_heading("1. Governing code")
    d.add_para(
        "HZISL uses the FIBA Official Basketball Rules 2024 as its baseline for the entire 2026–27 season, with the local modifications below. "
        "This prevents a mid-season rules change when FIBA’s 2026 edition takes effect on 1 October 2026. A later edition applies only if HZISL adopts it by written bulletin."
    )
    d.add_heading("2. Game and team format")
    d.add_table(
        ["Item", "HZISL rule"],
        [
            ["Players", "5 on court per team. A team must have 5 ready to start."],
            ["Matchday roster", "Maximum 12 eligible players, plus authorized team staff."],
            ["Playing time", "Four quarters of 8 minutes."],
            ["Intervals", "4 minutes between Q1–Q2 and Q3–Q4; 10 minutes at half-time."],
            ["Overtime", "3-minute periods with a 2-minute interval, repeated until a winner is determined."],
            ["Clock", "Stopped and restarted under FIBA principles; the referee may correct an obvious timing error before the report is signed."],
        ],
        [2050, 7310],
        body_size=9.5,
    )
    d.add_heading("3. Court, baskets and ball")
    d.add_bullets([
        "The host provides a safe court, securely fitted baskets, visible boundary lines, team benches and a controlled scorer’s table.",
        "Boys divisions use a size 7 ball; Girls divisions use a size 6 ball. The host supplies a match ball and a serviceable spare.",
        "The home team wears the lighter uniform and the away team the darker uniform unless both teams agree otherwise. Shirts must be numbered and distinguishable.",
        "Jewelry and unsafe equipment are not permitted. The referee decides whether protective equipment is safe."
    ])
    d.add_heading("4. Table officials and equipment")
    d.add_bullets([
        "The host appoints a trained scorer and timer. Their table record operates under the referee crew’s authority.",
        "The venue provides a visible game clock, score display, possession arrow, team-foul markers and audible signal where available.",
        "A 24/14-second shot clock is used only when compliant equipment and a trained operator are available and its use is declared to both teams before warm-up. Without it, deliberate delay may still be sanctioned under the governing rules.",
        "Instant replay and coach’s challenge procedures are not used in HZISL matches."
    ])

    d.add_heading("5. Start, possession and scoring")
    d.add_bullets([
        "The first quarter begins with a jump ball. Alternating possession applies thereafter in accordance with FIBA procedure.",
        "A successful free throw scores 1 point, a field goal from the 2-point area scores 2, and a field goal from the 3-point area scores 3.",
        "Teams change baskets for the second half. The scorer confirms the running score with the officials at each interval.",
        "Backcourt, closely guarded, throw-in, free-throw and restricted-area timing provisions follow the FIBA baseline, subject to the available table equipment."
    ])
    d.add_heading("6. Substitutions and time-outs")
    d.add_bullets([
        "There is no limit on the number of substitutions. Any of the 12 named players may enter during a substitution opportunity after the scorer’s signal and referee authorization.",
        "Each team may use 2 time-outs in the first half, 3 in the second half (no more than 2 in the final 2 minutes of Q4) and 1 in each overtime.",
        "Unused time-outs do not carry into the next half or overtime.",
        "An injured or bleeding player is managed under FIBA procedure and the school medical protocol. A suspected concussion or serious injury ends participation for the day unless appropriately cleared."
    ])
    d.add_heading("7. Violations and fouls")
    d.add_bullets([
        "Traveling, illegal dribble, out-of-bounds, 3-second, 5-second and backcourt violations follow the FIBA baseline.",
        "A player who commits 5 fouls is informed by an official and must leave the game immediately.",
        "Beginning with the fifth team foul in a quarter, non-shooting personal fouls are penalized under the FIBA team-foul rule.",
        "Technical, unsportsmanlike and disqualifying fouls use the FIBA 2024 terminology for this HZISL season.",
        "A disqualified player or team official leaves the playing area and is automatically suspended from the next HZISL match pending league review.",
    ])
    d.add_heading("8. Forfeit, suspension and safety")
    d.add_bullets([
        "A team unable to field 5 ready players within 15 minutes after the scheduled start normally forfeits 20–0, subject to Coordinator review of exceptional circumstances.",
        "A game may be suspended for unsafe court conditions, medical emergency, facility intrusion, severe behavior or another serious risk.",
        "The officials and host record the score, time, possession and reason for a suspension. The Coordinator decides whether the game resumes, is replayed, stands or is forfeited.",
    ])
    d.add_heading("9. Results and standings")
    d.add_table(
        ["Result", "League points"],
        [["Win", "2"], ["Loss after playing", "1"], ["Forfeit/default", "0"]],
        [6200, 3160],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    d.add_para(
        "Teams are ranked by league points, then head-to-head results among tied teams, head-to-head point difference, overall point difference, points scored and, if still required for a final placing, a neutral playoff or draw determined by the Council."
    )
    d.add_heading("10. Match report and protest")
    d.add_numbered([
        "The scorer and referee crew reconcile the running score, period totals, fouls and exceptional incidents.",
        "Both coaches review and sign the Official Match Report immediately after the game.",
        "A coach notes an intention to protest before signing when practicable; the school representative submits the written protest within 24 hours.",
        "A protest may concern eligibility, administration or a correctable rules application. It does not reopen ordinary judgment calls after the game."
    ])
    d.add_heading("Official references")
    d.add_source("FIBA Official Basketball Rules 2024", FIBA_2024_URL, "the HZISL baseline for the full 2026–27 season")
    d.add_source("FIBA rules resource hub", FIBA_RULES_HUB_URL, "official rules and interpretations")
    return d.save("HZISL_Basketball_Rules_2026-27.docx")


def build_volleyball_rules() -> Path:
    d = HZDocument(
        "Volleyball Competition Rules",
        "Volleyball Rules",
        "hzisl-sports-library/hzisl-volleyball.png",
        version=SPORT_RULES_VERSION,
    )
    d.add_cover(
        "Sport regulations",
        "Applies to MS Boys, MS Girls, HS Boys and HS Girls.",
    )
    d.add_callout(
        "Match format",
        "6-a-side • best of 3 sets • sets 1–2 to 25 • deciding set to 15 • win by 2 • unlimited substitutions.",
    )
    d.add_heading("1. Governing code")
    d.add_para(
        "HZISL uses the FIVB Official Volleyball Rules 2025–2028 as its baseline for the 2026–27 season, with the local modifications below. "
        "These local rules control where they differ from FIVB. The HZISL Competition Handbook controls league administration, eligibility, safety and protests."
    )
    d.add_heading("2. Match and team format")
    d.add_table(
        ["Item", "HZISL rule"],
        [
            ["Players", "6 on court per team. A team must have 6 eligible players ready to start."],
            ["Matchday roster", "Maximum 14 eligible players, including up to 2 designated liberos, plus authorized team staff."],
            ["Match format", "Best of 3 sets. The first team to win 2 sets wins the match."],
            ["Set scoring", "Sets 1 and 2 are played to 25 points; a deciding third set is played to 15. Every set must be won by 2 points, with no cap."],
            ["Intervals", "3 minutes between sets. The referee may extend an interval only for safety or facility necessity."],
            ["League tie", "A volleyball match cannot finish as a draw."],
        ],
        [2050, 7310],
        body_size=9.35,
    )
    d.add_heading("3. Court, net and ball")
    d.add_bullets([
        "The host provides a safe 18 × 9 metre court where practicable, visible boundary lines, safe free space, protected posts, team benches and a controlled scorer’s table.",
        "Net height is 2.43 m for Boys divisions and 2.24 m for Girls divisions unless HZISL publishes a division-specific safety adjustment before the season.",
        "The host supplies an approved indoor volleyball in good condition and at least one serviceable spare of the same type.",
        "The referee may require hazards to be removed, adjust a non-material local marking issue with both coaches informed, or refuse an unsafe court.",
    ])
    d.add_heading("4. Uniforms and equipment")
    d.add_bullets([
        "Teams wear matching numbered shirts and distinguishable colours. A libero wears a clearly contrasting shirt.",
        "Non-marking court footwear is required. Knee pads are strongly recommended.",
        "Jewellery and unsafe equipment are not permitted. Medical or religious coverings must be safe and approved at the referee’s inspection.",
        "If colours conflict, the home team changes or uses safe numbered bibs unless another solution was agreed in advance.",
    ])

    d.add_heading("5. Rotation, service and playing the ball")
    d.add_bullets([
        "Rally scoring is used. The team winning a rally scores 1 point and serves next.",
        "Six starting players rotate one position clockwise when their team wins the right to serve after receiving.",
        "A team may use up to 3 contacts to return the ball, in addition to a block contact. A player may not make two consecutive contacts except where FIVB rules permit it after a block or on one continuous action.",
        "The ball may touch the net while crossing on service or during a rally. A served ball must be contacted from behind the end line after the referee’s authorization.",
        "Position, rotation, service order, net-contact, centre-line, attack-line and back-row restrictions follow the FIVB baseline.",
    ])
    d.add_heading("6. Unlimited substitutions and libero")
    d.add_bullets([
        "There is no limit on the number of substitutions in a set or match. A substituted player may return later and may exchange with any eligible teammate.",
        "Substitutions occur only while the ball is out of play, after the scorer is ready and the referee authorizes entry. Teams must not use repeated requests to delay play.",
        "All substitutions are recorded by the scorer even though they are unlimited. A player may participate only if named on the certified matchday roster.",
        "Libero replacements are not counted as substitutions and otherwise follow FIVB libero procedure, including contrasting uniform and back-row restrictions.",
        "A suspected concussion or serious injury is managed under the school medical protocol; student welfare overrides competitive considerations.",
    ])
    d.add_heading("7. Time-outs, officials and conduct")
    d.add_bullets([
        "Each team may request 2 time-outs of 30 seconds in each set. Unused time-outs do not carry to another set.",
        "The first referee controls play. The second referee and scorer support substitutions, rotations, time-outs, sanctions and the official score.",
        "Delay warnings, delay penalties, misconduct sanctions and expulsion/disqualification follow the FIVB baseline. Serious conduct is also reported under the HZISL discipline process.",
        "There is no video challenge system in HZISL matches. Referee decisions on facts connected with play are final.",
    ])
    d.add_heading("8. Forfeit, suspension and safety")
    d.add_bullets([
        "A team unable to field 6 ready players within 15 minutes after the scheduled start normally forfeits 0–2, with sets recorded 0–25, subject to Coordinator review of exceptional circumstances.",
        "A team declared incomplete during play loses the set or match under the FIVB baseline, with the opponent receiving the points and sets needed to win.",
        "A match may be suspended for unsafe court conditions, medical emergency, facility intrusion, severe behaviour or another serious risk.",
        "The officials and host record the set score, point score and reason for a suspension. The Coordinator decides whether the match resumes, is replayed, stands or is forfeited.",
    ])

    d.add_heading("9. Results and standings", page_break_before=True)
    d.add_table(
        ["Result", "League points"],
        [["Win", "2"], ["Loss after playing", "1"], ["Forfeit/default", "0"]],
        [6200, 3160],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    d.add_para(
        "Teams are ranked by league points, then matches won, head-to-head result among tied teams, set difference, point difference, sets won, points scored and, if still required for a final placing, a neutral playoff or draw determined by the Council."
    )
    d.add_heading("10. Match report and protest")
    d.add_numbered([
        "The scorer and referee crew reconcile each set score, the final sets won, substitutions, sanctions and exceptional incidents.",
        "Both coaches review and sign the Official Match Report immediately after the match.",
        "A coach notes an intention to protest before signing when practicable; the school representative submits the written protest within 24 hours.",
        "A protest may concern eligibility, administration or a rules-application error. It does not reopen an ordinary referee judgement on facts connected with play.",
    ])
    d.add_heading("Official references")
    d.add_source("FIVB Official Volleyball Rules 2025–2028", FIVB_RULES_URL, "the international baseline for HZISL volleyball")
    d.add_source("FIVB basic rules", FIVB_BASIC_RULES_URL, "official overview of scoring, rotation and play")
    return d.save("HZISL_Volleyball_Rules_2026-27.docx")


def build_host_guide() -> Path:
    d = HZDocument("Host a Game Guide", "Host Guide")
    d.add_cover(
        "Matchday operations",
        "A practical timeline and checklist for safe, welcoming HZISL Saturdays.",
        compact=True,
    )
    d.add_callout(
        "Hosting standard",
        "The host owns venue readiness, visitor welcome, qualified officials, emergency response and submission of the signed result.",
    )
    d.add_heading("1. Assign the matchday team")
    d.add_table(
        ["Role", "Minimum responsibility"],
        [
            ["Event Lead", "Single point of contact; opens/closes venue; makes operational decisions."],
            ["Safeguarding lead", "Manages student supervision, changing-area boundaries and any safeguarding concern."],
            ["First-aid responder", "Holds first-aid equipment, emergency plan and access route for emergency services."],
            ["Officials coordinator", "Confirms referee crew and briefs table/field officials."],
            ["Scorer / timer", "Maintains the match record and supports immediate post-game reconciliation."],
            ["Facility marshal", "Controls access, spectator areas, water, toilets and team routes."],
        ],
        [2350, 7010],
        body_size=9.2,
    )
    d.add_heading("2. Preparation timeline")
    d.add_table(
        ["When", "Host action"],
        [
            ["14 days before", "Confirm venue booking, access, playing areas, emergency plan and school approval."],
            ["7 days before", "Confirm opponent contact, officials, first aid, equipment, changing arrangements and campus entry instructions."],
            ["72 hours before", "Send venue map, arrival route, parking/drop-off, surface notes, uniform colors and emergency contact."],
            ["24 hours before", "Recheck weather/air quality, rosters, officials, keys, clocks, balls, forms, water and cleaning."],
            ["06:30 if needed", "Issue an early safety/cancellation recommendation to opponent and Coordinator."],
            ["After the event", "Submit result and incidents; secure personal records; reset and close the venue."],
        ],
        [1800, 7560],
        body_size=9.3,
    )

    d.add_heading("3. Saturday run of show", page_break_before=True)
    d.add_table(
        ["Time", "Action", "Completion standard"],
        [
            ["07:15", "Open venue and verify access", "Gates, toilets, team routes and emergency access are open and safe."],
            ["07:30", "Prepare playing areas", "Field/court, goals/baskets, balls, clocks, benches, water and first aid ready."],
            ["07:45", "Welcome teams and officials", "Contacts confirmed; changing/warm-up areas explained; roster and colors checked."],
            ["09:00", "Start MS fixtures", "Officials and scoring table in position; no avoidable delay."],
            ["Between games", "Reset and reconcile", "Area checked; equipment reset; MS incidents noted; HS teams directed safely."],
            ["10:30", "Start HS fixtures", "Officials and scoring table ready; changes communicated."],
            ["Immediately after", "Confirm official record", "Final score, sanctions and incidents reviewed; coaches and official sign."],
            ["Within 30 min", "Send report to HZISL", "Clear image/file submitted through the approved channel."],
        ],
        [1150, 2880, 5330],
        body_size=8.8,
    )
    d.add_heading("4. Venue checklist")
    d.add_bullets([
        "Playing surface clear of hazards; goals anchored or baskets secure; boundaries visible.",
        "Warm-up and spectator areas separated from active play and emergency access.",
        "Team benches/technical areas, officials’ space and scorer’s table clearly identified.",
        "Drinking-water access, toilets and changing arrangements clean, open and explained.",
        "First-aid kit, ice/cold packs, charged phone, emergency contacts and ambulance access route ready.",
        "Blank Official Match Report and Incident Report available, with pens and a charged scanning/photograph device.",
    ])
    d.add_heading("5. Sport-specific equipment", page_break_before=True)
    d.add_table(
        ["Football", "Basketball"],
        [
            ["Size 5 match ball plus at least 2 spares", "Correct ball: size 7 Boys / size 6 Girls, plus spare"],
            ["Anchored goals, nets, corner markers, team benches", "Game clock, score display, audible signal, possession arrow"],
            ["Assistant-referee flags where crew uses them", "Team-foul markers; shot clock only if trained operator is confirmed"],
            ["Safe technical and spectator boundaries", "Safe scorer’s table and clear bench boundaries"],
        ],
        [4680, 4680],
        body_size=9.0,
    )
    d.add_heading("6. Visitor welcome and briefing")
    d.add_numbered([
        "Meet the visiting coach and official crew at the agreed point.",
        "Confirm the Event Lead, first-aid responder and emergency contact.",
        "Show toilets, changing arrangements, warm-up area, water and spectator zone.",
        "Confirm uniform colors, match format and any declared equipment limitation.",
        "Give the official the certified matchday rosters and explain the result-signing process.",
    ])

    d.add_heading("7. Safety and incident response")
    d.add_callout(
        "Stop first",
        "For a serious injury, suspected concussion, safeguarding concern, lightning, unsafe facility or threatening behavior: stop the activity, protect the student, activate the host plan and notify the appropriate school leaders.",
        fill=PALE_RED,
        accent=RED,
    )
    d.add_bullets([
        "The responsible school adult contacts the student’s parent/guardian under school procedure.",
        "Call emergency services when indicated; do not delay emergency care to complete league paperwork.",
        "Preserve privacy. Record objective facts, actions and times; do not diagnose, speculate or share sensitive details in group chats.",
        "Use the Incident, Injury & Disciplinary Report for any serious injury, ejection, abuse, facility danger, emergency-service response or safeguarding escalation.",
        "Completed sensitive forms go only through the approved secure channel and are not uploaded to the public app.",
    ])
    d.add_heading("8. Result close-out")
    d.add_numbered([
        "Scorer and official reconcile the final score and all period/half totals.",
        "Record cautions, sendings-off, technical/disqualifying fouls, injuries and protest notice.",
        "Home coach, away coach, lead official and host scorer sign the Official Match Report.",
        "Photograph or scan the complete signed page so every field is legible.",
        "Send it within 30 minutes. Keep the original securely until the Coordinator confirms receipt.",
        "Close spectator and team areas, check for lost property/damage, reset equipment and lock the venue."
    ])
    d.add_heading("9. Handover message template")
    d.add_callout(
        "Send",
        "Matchday [#] • [Division] • [Home] [score]–[score] [Away] • report attached • incidents: none / Incident Report attached • submitted by [name] at [time].",
        fill=PALE_GOLD,
        accent=RED,
    )
    return d.save("HZISL_Host_a_Game_Guide_2026-27.docx")


def _field_table(d: HZDocument, rows: Sequence[tuple[str, str]], *, widths=(2600, 6760),
                 body_size: float = 9.5) -> None:
    d.add_table(["Field", "Response"], rows, list(widths), body_size=body_size, header_fill=PALE_GRAY)


def _signature_table(d: HZDocument, roles: Sequence[str], *, blank_lines: int = 2) -> None:
    table = d.doc.add_table(rows=1, cols=3)
    widths = [2600, 4200, 2560]
    _style_table_borders(table)
    for idx, header in enumerate(("Role", "Printed name and signature", "Date / time")):
        _set_cell_text(table.rows[0].cells[idx], header, size=8.7, bold=True, color=NAVY, fill=PALE_BLUE)
    _set_repeat_header(table.rows[0])
    for role in roles:
        row = table.add_row()
        _set_cell_text(row.cells[0], role, size=8.9, bold=True)
        blank = "\n" * blank_lines
        _set_cell_text(row.cells[1], blank, size=8.9)
        _set_cell_text(row.cells[2], blank, size=8.9)
    _set_table_geometry(table, widths)


def build_match_report() -> Path:
    d = HZDocument("Official Match Report", "Match Report")
    d.add_cover(
        "Complete immediately after the game",
        "The signed form is the official source for the result, sanctions and any protest notice.",
        compact=True,
    )
    d.add_callout(
        "Submission",
        "Host sends a clear scan/photo to the HZISL Coordinator within 30 minutes. Keep the original until receipt is confirmed. A signature confirms review; it does not waive a timely protest.",
        fill=PALE_GOLD,
        accent=RED,
    )
    _field_table(d, [
        ("Match", "Date: __________________  Matchday: ______  Scheduled start: ______  Actual finish: ______"),
        ("Competition", "□ Football  □ Basketball  □ Volleyball     □ MS  □ HS     □ Boys  □ Girls"),
        ("Venue", "____________________________________________________________________________"),
        ("Teams", "Home: ____________________________________   Away: ____________________________________"),
        ("Officials", "Lead referee: ______________________________   Other official(s): ___________________________"),
    ], body_size=9.0)
    d.add_heading("Official score", level=2)
    d.add_table(
        ["Team", "P1 / H1", "P2 / H2", "P3", "P4", "OT", "FINAL"],
        [
            ["HOME", "", "", "", "", "", ""],
            ["AWAY", "", "", "", "", "", ""],
        ],
        [2500, 1050, 1050, 1050, 1050, 1050, 1610],
        body_size=10,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6,
    )
    d.add_para(
        "Football: use H1/H2 and FINAL. Basketball: use P1–P4, OT if needed, and FINAL. Volleyball: use set columns and FINAL.",
        size=8.5,
        color=MUTED,
        italic=True,
        after=4,
    )
    d.add_heading("Sanctions and incidents", level=2, page_break_before=True)
    _field_table(d, [
        ("Match identification", "Date: __________________  Home: __________________  Away: __________________"),
    ], body_size=8.8)
    d.add_table(
        ["Team", "Player / staff", "No.", "Time / period", "Code and brief reason"],
        [["", "", "", "", ""], ["", "", "", "", ""], ["", "", "", "", ""]],
        [1100, 2350, 700, 1350, 3860],
        body_size=8.8,
    )
    d.add_para(
        "Codes: Football YC / RC. Basketball TF / UF / DQF. Serious injury, emergency response or major conduct issue requires a separate Incident Report.",
        size=8.4,
        color=MUTED,
        after=4,
    )
    d.add_table(
        ["Post-game declaration", "Mark and explain"],
        [
            ["Result status", "□ Confirmed   □ Suspended/abandoned   □ Forfeit requested"],
            ["Injury / incident", "□ None requiring report   □ Incident Report attached"],
            ["Protest notice", "□ None   □ Home intends to protest   □ Away intends to protest\nRule/issue noted: ________________________________________________"],
        ],
        [2600, 6760],
        body_size=8.8,
        header_fill=PALE_GRAY,
    )
    d.add_heading("Signatures", level=2)
    _signature_table(d, ["Home head coach", "Away head coach", "Lead official", "Host scorer / Event Lead"])
    return d.save("HZISL_Official_Match_Report_2026-27.docx")


def build_agreement() -> Path:
    d = HZDocument("School Participation Agreement", "Participation Agreement")
    d.add_cover(
        "Annual league commitment",
        "For completion by each HZISL member school for the 2026–27 season.",
        compact=True,
    )
    _field_table(d, [
        ("Member school", "____________________________________________________________________________"),
        ("School representative", "Name: __________________________  Role: __________________________"),
        ("Contact", "Email: ______________________________  Mobile: _________________________"),
        ("Participating teams", "□ Football MS Boys  □ MS Girls  □ HS Boys  □ HS Girls\n□ Basketball MS Boys  □ MS Girls  □ HS Boys  □ HS Girls\n□ Volleyball MS Boys  □ MS Girls  □ HS Boys  □ HS Girls"),
    ])
    d.add_heading("1. Purpose")
    d.add_para(
        "The school joins HZISL as a cooperative member to provide safe, fair and dependable interschool competition. "
        "This Agreement is read with the Competition Handbook, sport rules, host guide and official forms."
    )
    d.add_heading("2. School commitments")
    d.add_bullets([
        "Appoint a representative authorized to receive notices and coordinate teams.",
        "Certify enrollment, division eligibility, consent and medical clearance for every participant.",
        "Provide qualified adult supervision and follow the school’s safeguarding and emergency duties at home and away fixtures.",
        "Honor the published schedule, communicate conflicts early and never treat an unconfirmed request as a schedule change.",
        "When hosting, provide a safe venue, first aid, officials, match administration and timely submission of signed results.",
        "Ensure players, coaches and spectators follow the HZISL conduct standards and cooperate with fair discipline review.",
        "Protect personal information and use secure channels for completed rosters and incident records.",
    ], size=10.3, after=2)
    d.add_heading("3. HZISL commitments")
    d.add_bullets([
        "Publish the schedule, competition documents, contacts and verified public results.",
        "Apply rules consistently, manage declared conflicts of interest and give schools a fair chance to respond to material concerns.",
        "Limit collected information to what the competition reasonably needs and keep sensitive records out of the public app.",
        "Coordinate Council decisions, protests, rescheduling and season certification in a timely manner.",
    ], size=10.3, after=2)
    d.add_heading("4. Risk, insurance and authority")
    d.add_para(
        "Each school remains responsible for its students, staff, insurance, permissions, medical processes and premises. "
        "HZISL coordination does not replace a school’s legal, safeguarding or duty-of-care obligations. The host may close a facility and officials may stop play when safety requires it."
    )
    d.add_heading("5. Term, withdrawal and disputes")
    d.add_bullets([
        "This Agreement covers the 2026–27 HZISL season and ends after final records and outstanding matters are closed.",
        "A school intending to withdraw gives as much notice as possible, normally at least 30 days. The Council determines schedule and standing consequences.",
        "Operational disputes first go to the Coordinator. A material unresolved dispute is referred to neutral Council representatives under the Handbook.",
        "Any approved participation fee or shared cost must be documented separately; this Agreement alone creates no unspecified payment obligation.",
    ], size=10.3, after=2)
    d.add_heading("6. Acceptance")
    d.add_para(
        "By signing, the school confirms that it has reviewed the current HZISL documents, will brief its teams and has authority to make this seasonal commitment."
    )
    _signature_table(
        d,
        ["School Head / authorized signatory", "School HZISL representative", "HZISL Coordinator"],
        blank_lines=1,
    )
    return d.save("HZISL_School_Participation_Agreement_2026-27.docx")


def build_roster() -> Path:
    d = HZDocument("Team Roster & Eligibility Declaration", "Roster & Eligibility")
    d.add_cover(
        "Secure administrative form",
        "Complete one form per team and submit before the team’s first fixture.",
        compact=True,
    )
    d.add_callout(
        "Privacy",
        "Send completed forms only through the approved secure channel. Do not upload student names, identifiers or medical information to the public app.",
        fill=PALE_RED,
        accent=RED,
    )
    _field_table(d, [
        ("School / team", "School: ______________________________  Team name: ______________________________"),
        ("Competition", "□ Football  □ Basketball  □ Volleyball     □ MS  □ HS     □ Boys  □ Girls"),
        ("Staff", "Head coach: __________________________  School representative: _______________________"),
        ("Contacts", "Coach mobile: ________________________  Representative email: ________________________"),
    ], body_size=9.0)
    d.add_heading("Certified season roster", level=2)
    headers = ["#", "Student name", "School ID / initials", "Grade", "Jersey", "Eligible ✓"]
    roster_widths = [620, 3100, 2000, 1000, 1000, 1640]
    roster_alignments = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    d.add_table(
        headers,
        [[str(i), "", "", "", "", "□"] for i in range(1, 11)],
        roster_widths,
        body_size=8.3,
        alignments=roster_alignments,
    )
    d.add_heading("Certified season roster — continued", level=2, page_break_before=True)
    d.add_table(
        headers,
        [[str(i), "", "", "", "", "□"] for i in range(11, 21)],
        roster_widths,
        body_size=8.3,
        alignments=roster_alignments,
    )
    d.add_heading("School declaration", level=2)
    d.add_para(
        "I certify on behalf of the school that every listed student is currently enrolled, belongs to or is approved to play up into the declared division, "
        "has required consent and medical clearance, and is not listed for another school or conflicting HZISL team. The school retains emergency and medical records and will report roster changes before participation."
    )
    _signature_table(d, ["Head coach", "School HZISL representative / Athletic Director"])
    return d.save("HZISL_Team_Roster_and_Eligibility_2026-27.docx")


def build_incident_report() -> Path:
    d = HZDocument("Incident, Injury & Disciplinary Report", "Incident Report")
    d.add_cover(
        "Confidential follow-up form",
        "Use for serious injury, ejection, major conduct, facility danger, emergency response or safeguarding escalation.",
        compact=True,
    )
    d.add_callout(
        "Emergency first",
        "Protect the student and activate the host/school emergency process before completing this form. Safeguarding concerns follow the required reporting route immediately.",
        fill=PALE_RED,
        accent=RED,
    )
    _field_table(d, [
        ("Match", "Date: __________________  Matchday: ______  Start time: ______  Incident time: ______"),
        ("Competition", "□ Football  □ Basketball  □ Volleyball     □ MS  □ HS     □ Boys  □ Girls"),
        ("Venue / teams", "Venue: __________________________  Home: __________________  Away: __________________"),
        ("Report type", "□ Injury  □ Ejection/discipline  □ Conduct  □ Facility/safety  □ Safeguarding  □ Other"),
        ("Reporter", "Name: __________________________  Role/school: __________________  Mobile: ______________"),
    ], body_size=8.9)
    d.add_heading("1. Person(s) directly involved", level=2)
    d.add_table(
        ["Name / identifier", "Role and school", "Age/grade if student", "Parent/guardian notified"],
        [["", "", "", "□ Yes  □ No  □ N/A"], ["", "", "", "□ Yes  □ No  □ N/A"]],
        [2600, 2600, 1900, 2260],
        body_size=8.6,
    )
    d.add_heading("2. Objective account", level=2)
    d.add_table(
        ["Describe what was seen/heard, where it happened and the sequence of events. Avoid diagnosis or speculation."],
        [["\n\n\n\n\n"]],
        [9360],
        body_size=9.0,
        header_fill=PALE_GRAY,
    )
    d.add_heading("3. Immediate action", level=2)
    d.add_table(
        ["Action", "Details / time"],
        [
            ["Play stopped / area secured", "□ Yes  □ No  Time: ______  Details: ______________________________"],
            ["First aid / medical assessment", "□ Yes  □ No  By whom: _________________________________________"],
            ["Emergency services", "□ Called  □ Not called  Time: ______  Outcome/destination: ________________"],
            ["School leaders / safeguarding lead", "Who was notified and when: __________________________________________"],
            ["Participant removed / ejected", "Person and destination/supervising adult: ______________________________"],
        ],
        [3100, 6260],
        body_size=8.6,
        header_fill=PALE_GRAY,
    )
    d.add_heading("4. Officials, witnesses and evidence")
    d.add_table(
        ["Name", "Role / school", "Contact held by", "Statement requested"],
        [["", "", "", "□"], ["", "", "", "□"], ["", "", "", "□"]],
        [2500, 2400, 2500, 1960],
        body_size=8.8,
    )
    d.add_para(
        "Evidence preserved (do not attach sensitive media to public/group channels):\n"
        "□ Official Match Report  □ Official’s statement  □ Witness statement  □ Photo/video  □ Medical note held by school  □ Other: __________________",
        size=9.2,
    )
    d.add_heading("5. Follow-up requested")
    d.add_table(
        ["Area", "Request / owner / deadline"],
        [
            ["Student welfare", "________________________________________________________________"],
            ["Disciplinary review", "________________________________________________________________"],
            ["Facility correction", "________________________________________________________________"],
            ["Communication", "________________________________________________________________"],
        ],
        [2600, 6760],
        body_size=9.0,
    )
    d.add_callout(
        "Secure submission",
        "Send to the HZISL Coordinator and the required school authority through the approved secure channel as soon as practicable. Never publish a completed incident form in the HZISL app.",
        fill=PALE_GOLD,
        accent=RED,
    )
    d.add_heading("6. Reporter certification")
    d.add_para(
        "I have recorded the information honestly and as objectively as possible. I understand that safeguarding and medical records must be handled under school policy and applicable law, and that this form is shared only with people who need it for safety or league administration."
    )
    _signature_table(
        d,
        ["Person completing report", "Host Event Lead", "School safeguarding/athletics lead (if required)"],
        blank_lines=1,
    )
    return d.save("HZISL_Incident_Injury_Disciplinary_Report_2026-27.docx")


def main() -> None:
    builders = {
        "handbook": build_handbook,
        "football-rules": build_football_rules,
        "basketball-rules": build_basketball_rules,
        "volleyball-rules": build_volleyball_rules,
        "host-guide": build_host_guide,
        "match-report": build_match_report,
        "agreement": build_agreement,
        "roster": build_roster,
        "incident-report": build_incident_report,
    }
    parser = argparse.ArgumentParser(description="Build HZISL competition DOCX files.")
    parser.add_argument(
        "--only",
        nargs="+",
        choices=tuple(builders),
        help="Build only the named documents. Without this option, every document is built.",
    )
    args = parser.parse_args()
    selected = args.only or list(builders)
    paths = [builders[name]() for name in selected]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
